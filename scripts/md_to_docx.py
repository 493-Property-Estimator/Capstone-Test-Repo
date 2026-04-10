#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

import markdown as md
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def sanitize_text(text: str) -> str:
    # Remove emoji-like ranges and presentation selectors that sometimes render as emojis in Word.
    text = "".join(
        ch
        for ch in text
        if not (
            (0x1F300 <= ord(ch) <= 0x1FAFF)
            or (0x2600 <= ord(ch) <= 0x26FF)
            or (0x2700 <= ord(ch) <= 0x27BF)
            or ord(ch) in (0xFE0F, 0x200D)
        )
    )
    # Replace a few common symbols with ASCII so they can't render as icons.
    return (
        text.replace("↔", "<->")
        .replace("–", "-")
        .replace("—", "-")
    )


def set_document_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.05
    style.paragraph_format.space_after = Pt(3)

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_code_block(document: Document, code_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["No Spacing"] if "No Spacing" in document.styles else document.styles["Normal"]
    run = paragraph.add_run(sanitize_text(code_text.rstrip("\n")))
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def add_list_item(document: Document, text: str, ordered: bool) -> None:
    text = sanitize_text(text)
    style_name = "List Number" if ordered else "List Bullet"
    if style_name in document.styles:
        document.add_paragraph(text, style=style_name)
    else:
        prefix = "1. " if ordered else "- "
        document.add_paragraph(prefix + text)


def add_table(document: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return
    cell_matrix: list[list[str]] = []
    for row in rows:
        cells = row.find_all(["th", "td"])
        cell_matrix.append([sanitize_text(cell.get_text(" ", strip=True)) for cell in cells])
    column_count = max((len(r) for r in cell_matrix), default=0)
    if column_count == 0:
        return

    table = document.add_table(rows=len(cell_matrix), cols=column_count)
    table.style = "Table Grid" if "Table Grid" in document.styles else table.style
    for row_index, row_cells in enumerate(cell_matrix):
        for col_index in range(column_count):
            cell_text = row_cells[col_index] if col_index < len(row_cells) else ""
            table.cell(row_index, col_index).text = cell_text


def add_heading(document: Document, text: str, level: int) -> None:
    text = sanitize_text(text).strip()
    if not text:
        return
    level = max(1, min(9, level))
    document.add_heading(text, level=level)


def is_block(tag: Tag) -> bool:
    return tag.name in {
        "p",
        "pre",
        "ul",
        "ol",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "hr",
        "table",
        "blockquote",
    }


def render_html_to_docx(document: Document, soup: BeautifulSoup) -> None:
    body = soup.body or soup
    for element in body.descendants:
        if isinstance(element, NavigableString):
            continue
        if not isinstance(element, Tag):
            continue
        if not is_block(element):
            continue

        if element.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(element.name[1])
            add_heading(document, element.get_text(" ", strip=True), level=level)
            continue

        if element.name == "p":
            text = sanitize_text(element.get_text(" ", strip=True))
            if text:
                document.add_paragraph(text)
            continue

        if element.name in {"ul", "ol"}:
            ordered = element.name == "ol"
            for li in element.find_all("li", recursive=False):
                li_text = li.get_text(" ", strip=True)
                if li_text:
                    add_list_item(document, li_text, ordered=ordered)
            continue

        if element.name == "pre":
            code = element.get_text("", strip=False)
            add_code_block(document, code_text=code)
            continue

        if element.name == "table":
            add_table(document, element)
            continue

        if element.name == "blockquote":
            quote_text = sanitize_text(element.get_text("\n", strip=True))
            if quote_text:
                paragraph = document.add_paragraph(quote_text)
                paragraph.style = document.styles["Quote"] if "Quote" in document.styles else paragraph.style
            continue

        if element.name == "hr":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("-" * 30)
            run.font.size = Pt(10)
            continue


def convert_markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    html = md.markdown(
        markdown_text,
        extensions=["extra", "fenced_code", "tables", "sane_lists"],
        output_format="html5",
    )
    soup = BeautifulSoup(f"<html><body>{html}</body></html>", "lxml")

    document = Document()
    set_document_defaults(document)
    render_html_to_docx(document, soup)
    document.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert a Markdown file into a Word .docx file.")
    parser.add_argument("input_md", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()

    markdown_text = args.input_md.read_text(encoding="utf-8")
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    convert_markdown_to_docx(markdown_text, args.output_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
